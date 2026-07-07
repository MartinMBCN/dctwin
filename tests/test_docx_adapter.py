from __future__ import annotations

from pathlib import Path

from docx import Document

from dctwin.adapters import AdapterRegistry, DocxCvAdapter, PdfCvAdapter


def test_registry_selects_docx_strategy_and_extracts_paragraphs_and_tables(
    tmp_path: Path,
) -> None:
    path = tmp_path / "cv.docx"
    document = Document()
    document.add_paragraph("Alex Example | alex@example.com")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "ACME"
    table.cell(0, 1).text = "Platform Lead"
    document.save(path)

    registry = AdapterRegistry()
    registry.register(PdfCvAdapter())
    registry.register(DocxCvAdapter())
    result = registry.adapt("cv", path)
    adapted = result.model_document

    assert adapted["adapter"]["name"] == "docx_cv"
    assert adapted["source_type"] == "cv"
    assert adapted["blocks"][0]["text"] == "Alex Example | [EMAIL_1]"
    assert any(block["locator"]["kind"] == "record" for block in adapted["blocks"])
    assert result.enrollment_document["candidates"] == [
        {
            "type": "email",
            "value": "alex@example.com",
            "source_block_id": "block_paragraph_1",
            "purpose": "account_enrollment",
            "verification_status": "unverified",
        }
    ]
    assert "alex@example.com" not in repr(result)
