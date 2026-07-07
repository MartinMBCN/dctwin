from __future__ import annotations

import hashlib
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document

from dctwin.adapters.base import AdaptedSource
from dctwin.privacy import minimize_direct_identifiers


class DocxCvAdapter:
    """CV strategy for Word Open XML documents, including tables."""

    source_type = "cv"
    name = "docx_cv"
    version = "0.1.0"

    def supports(self, path: Path) -> bool:
        return path.is_file() and path.suffix.lower() == ".docx"

    def adapt(self, path: Path) -> AdaptedSource:
        raw_bytes = path.read_bytes()
        digest = hashlib.sha256(raw_bytes).hexdigest()
        source_id = f"src_{digest[:16]}"
        document = Document(path)
        blocks: list[dict[str, Any]] = []
        enrollment_candidates: list[dict[str, str]] = []
        redaction_counters: dict[str, int] = {}

        for paragraph_number, paragraph in enumerate(document.paragraphs, start=1):
            if paragraph.text.strip():
                block, candidates = self._block(
                    block_id=f"block_paragraph_{paragraph_number}",
                    locator_kind="paragraph",
                    locator_value=str(paragraph_number),
                    text=paragraph.text,
                    counters=redaction_counters,
                )
                blocks.append(block)
                enrollment_candidates.extend(candidates)

        for table_number, table in enumerate(document.tables, start=1):
            for row_number, row in enumerate(table.rows, start=1):
                text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if text:
                    block, candidates = self._block(
                        block_id=f"block_table_{table_number}_row_{row_number}",
                        locator_kind="record",
                        locator_value=f"table:{table_number}/row:{row_number}",
                        text=text,
                        counters=redaction_counters,
                    )
                    blocks.append(block)
                    enrollment_candidates.extend(candidates)

        if not blocks:
            raise ValueError(
                "The DOCX contains no extractable paragraph or table text; "
                "text boxes and embedded images require an additional strategy"
            )

        model_document = {
            "schema_version": "0.1.0",
            "source_id": source_id,
            "source_type": self.source_type,
            "media_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "content_hash": f"sha256:{digest}",
            "ingested_at": datetime.now(UTC).isoformat(),
            "adapter": {"name": self.name, "version": self.version},
            "privacy": {
                "classification": "personal_data",
                "direct_identifiers_minimized": True,
                "enrollment_candidates_separated": True,
            },
            "blocks": blocks,
        }
        return AdaptedSource(
            model_document=model_document,
            enrollment_document={
                "schema_version": "0.1.0",
                "source_id": source_id,
                "candidates": enrollment_candidates,
            },
        )

    @staticmethod
    def _block(
        *,
        block_id: str,
        locator_kind: str,
        locator_value: str,
        text: str,
        counters: dict[str, int],
    ) -> tuple[dict[str, Any], list[dict[str, str]]]:
        minimized, redactions = minimize_direct_identifiers(text, counters=counters)
        block = {
            "id": block_id,
            "locator": {"kind": locator_kind, "value": locator_value},
            "text": minimized.strip(),
            "redactions": [
                {"category": item.category, "placeholder": item.placeholder}
                for item in redactions
            ],
        }
        candidates = [
            {
                "type": "email",
                "value": item.value,
                "source_block_id": block_id,
                "purpose": "account_enrollment",
                "verification_status": "unverified",
            }
            for item in redactions
            if item.category == "email"
        ]
        return block, candidates
